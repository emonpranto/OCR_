import os
import pandas as pd
import pdfplumber
import streamlit as st
import uuid

PDF_STORAGE_PATH = 'document_store/pdfs/'

if not os.path.exists(PDF_STORAGE_PATH):
    os.makedirs(PDF_STORAGE_PATH)
    
class DocumentChunk:
    def __init__(self, section_title, content, metadata=None):
        self.id = str(uuid.uuid4())
        self.section_title = section_title
        self.page_content = content
        self.metadata = metadata or {}

# Save uploaded file
def save_uploaded_file(uploaded_file):
    file_path = os.path.join(PDF_STORAGE_PATH, uploaded_file.name)
    with open(file_path, "wb") as file:
        file.write(uploaded_file.getbuffer())
    return file_path

# Extracting text from pdf
def extract_text_from_pdf(file_path):
    try:
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n\n"
            return text.strip()
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return ""

# Extract tables from PDF
def extract_tables_from_pdf(file_path):
    try:
        with pdfplumber.open(file_path) as pdf:
            tables = []
            for page in pdf.pages:
                page_tables = page.extract_tables()
                if page_tables:
                    tables.extend(page_tables)
            return tables
    except Exception as e:
        st.error(f"Error extracting tables from PDF: {e}")
        return []

# Convert tables to DataFrames
def convert_tables_to_dataframe(tables):
    dfs = []
    for table in tables:
        if table and len(table) > 1:  # Ensure table has headers and data
            try:
                df = pd.DataFrame(table[1:], columns=table[0])
                dfs.append(df)
            except Exception as e:
                st.warning(f"Error converting table to DataFrame: {e}")
    return dfs


def load_table_file(file):
    """
    Load CSV or Excel file from uploaded file-like object.
    Returns a cleaned DataFrame.
    """
    try:
        if file.name.endswith('.csv'):
            df = pd.read_csv(file)
        else:  # Excel
            df = pd.read_excel(file)
    except Exception as e:
        st.error(f"Error loading table file: {e}")
        return None

    # Clean the dataframe
    df_clean = clean_dataframe(df)
    return df_clean

def clean_dataframe(df):
    """
    Clean and normalize the dataframe.
    """
    df = df.dropna(how='all').dropna(axis=1, how='all')  # Drop empty rows and columns

    # Strip whitespace from string columns
    for col in df.select_dtypes(include=['object']).columns:
        df[col] = df[col].str.strip()

    return df


from docx import Document
import pandas as pd

def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = ""
        for para in doc.paragraphs:
            full_text += para.text + "\n\n"
        return full_text.strip()
    except Exception as e:
        return f"Error extracting text: {e}"

# Extract tables from DOCX
def extract_tables_from_docx(file_path):
    try:
        doc = Document(file_path)
        tables = []
        for table in doc.tables:
            data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                data.append(row_data)
            if len(data) > 1:
                df = pd.DataFrame(data[1:], columns=data[0])
            else:
                df = pd.DataFrame(data)
            tables.append(df)
        return tables
    except Exception as e:
        st.error(f"Error extracting tables from DOCX: {e}")
        return []